import docx

document = docx.Document('list test.docx')
#print(document.paragraphs[0].style.name)

my_para = document.paragraphs[0]
my_para.style = 'ListNumber'

document.save('restyled test list.docx')
