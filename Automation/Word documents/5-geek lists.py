# Import docx NOT python-docx 
import docx 

# Create an instance of a word document 
doc = docx.Document() 

# Add a Title to the document 
doc.add_heading('GeeksForGeeks', 0) 

# Adding list of style name 'List Number' 
doc.add_heading('Style: List Number', 3) 
# Adding points to the list named 'List Number' 
doc.add_paragraph('The first item in an ordered list.', 
				style='List Number') 

doc.add_paragraph('The second item in an ordered list.', 
				style='List Number') 

doc.add_paragraph('The third item in an ordered list.', 
				style='List Number') 

# Adding list of style name 'List Number 2' 
doc.add_heading('Style: List Number 2', 3) 
# Adding points to the list named 'List Number 2' 
doc.add_paragraph('The first item in an ordered list.', 
				style='List Number 2') 

doc.add_paragraph('The second item in an ordered list.', 
				style='List Number 2') 

doc.add_paragraph('The third item in an ordered list.', 
				style='List Number 2') 

# Adding list of style name 'List Number 3' 
doc.add_heading('Style: List Number 3', 3) 
# Adding points to the list named 'List Number 3' 
doc.add_paragraph('The first item in an ordered list.', 
				style='List Number 3') 

doc.add_paragraph('The second item in an ordered list.', 
				style='List Number 3') 

doc.add_paragraph('The third item in an ordered list.', 
				style='List Number 3') 

# Now save the document to a location 
doc.save('gfg.docx')
