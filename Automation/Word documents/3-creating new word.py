# Import docx NOT python-docx 
import docx 
from docx.shared import Pt 

# Create an instance of a word document 
doc = docx.Document() 

# Add a Title to the document 
doc.add_heading('GeeksForGeeks', 0) 

# Adding paragraph with Increased font size 
doc.add_heading('Increased Font Size Paragraph:', 3) 
para = doc.add_paragraph().add_run( 
	'GeeksforGeeks is a Computer Science portal for geeks.') 
# Increasing size of the font 
para.font.size = Pt(12) 

# Adding paragraph with normal font size 
doc.add_heading('Normal Font Size Paragraph:', 3) 
doc.add_paragraph( 
	'GeeksforGeeks is a Computer Science portal for geeks.') 

# Now save the document to a location 
doc.save('gfg.docx')
