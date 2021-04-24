import docx

document = docx.Document('gfg.docx')

my_para = document.paragraphs[0]
list_para = document.paragraphs[2]
print(list_para.text)
print(list_para.style.name)

list_para.style = document.styles['List Bullet']

document.save('restyled gfg.docx')
