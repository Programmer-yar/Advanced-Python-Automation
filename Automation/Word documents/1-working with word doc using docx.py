import docx

doc = docx.Document('word test.docx')
#print(len(doc.paragraphs))

print(doc.paragraphs[2].runs[0].text)
print("total number of runs of first list element\n", doc.paragraphs[2].runs)
#doc.paragraphs[2].style = 'Normal'
#style object has 3 attributes: type, name, style_id
print('first para:' ,doc.paragraphs[0].text)
print('first para style: ', doc.paragraphs[0].style)

print('first list style', doc.paragraphs[3].style)
print(doc.paragraphs[3].runs[0].style)


print(doc.paragraphs[3].style.name)

print(doc.paragraphs[8].style)
